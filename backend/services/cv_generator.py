from docx import Document
from fpdf import FPDF
import os
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

class CVGenerator:
    @staticmethod
    def generate_ats_docx(cv_data, output_path=None):
        doc = Document()
        
        # Adjust top margin
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)  # Reduce top margin to 0.5 inches

        # Add Name centered
        name_heading = doc.add_heading(cv_data['name'], level=1)
        name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Reduce space after name heading
        name_heading_format = name_heading.paragraph_format
        name_heading_format.space_after = Pt(12)  # Reduce space after heading to 12pt

        # Format contact info in one line
        contact_info = cv_data['contact_info']
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if isinstance(contact_info, dict):
            contact_parts = []
            if contact_info.get('phone_number'):
                contact_parts.append(f"Tel: {contact_info['phone_number']}")
            if contact_info.get('email'):
                contact_parts.append(f"Email: {contact_info['email']}")
            if contact_info.get('location'):
                contact_parts.append(f"Address: {contact_info['location']}")
            contact_paragraph.add_run(' | '.join(contact_parts))

        # Add Professional Summary with heading
        prof_heading = doc.add_heading('PROFESSIONAL PROFILE', level=2)
        prof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(cv_data['professional_summary'])

        # Add Skills section
        skills_heading = doc.add_heading('SKILLS', level=2)
        skills_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Format skills with vertical bars
        skills_text = ' | '.join(cv_data['skills'])
        doc.add_paragraph(skills_text)

        # Add Key Achievements
        if cv_data.get('key_achievements'):
            achieve_heading = doc.add_heading('KEY ACHIEVEMENTS', level=2)
            achieve_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for achievement in cv_data['key_achievements']:
                doc.add_paragraph(f"{achievement}", style='List Bullet')

        # Add Experience section
        exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        exp_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for job in cv_data['experience']:
            # Add job title and company on one line
            job_heading = doc.add_paragraph()
            job_heading.add_run(f"{job['job_title']} | {job['date_range']}").bold = True
            if job.get('company'):
                doc.add_paragraph(job['company'])
            # Add bullet points
            for bullet in job['bullet_points']:
                doc.add_paragraph(f"{bullet}", style='List Bullet')

        # Add Education with centered heading
        edu_heading = doc.add_heading('EDUCATION', level=2)
        edu_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for edu in cv_data['education']:
            doc.add_paragraph(
                f"{edu['degree']} - {edu['institution']} ({edu['graduation_year']})"
            )

        # Add Certifications with centered heading
        if cv_data.get('certifications'):
            cert_heading = doc.add_heading('CERTIFICATIONS', level=2)
            cert_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cert in cv_data['certifications']:
                doc.add_paragraph(
                    f"{cert['name']} ({cert['year']}) - {cert['issuing_organization']}"
                )

        if output_path:
            doc.save(output_path)
            return output_path
        
        return doc

    @staticmethod
    def generate_modern_pdf(cv_data, output_path):
        pdf = FPDF()
        pdf.add_page()
        
        # Set default font
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf')
        pdf.add_font('DejaVu', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf')
        
        # Add Name and Contact Info
        pdf.set_font('DejaVu', 'B', 16)
        pdf.write_html(f'<h1 style="text-align: center;">{cv_data["name"]}</h1>')
        pdf.set_font('DejaVu', '', 12)
        pdf.write_html(f'<p style="text-align: center;">{cv_data["contact_info"]}</p>')
        pdf.ln(10)

        # Add Professional Summary
        pdf.set_font('DejaVu', 'B', 14)
        pdf.write_html('<h2>Professional Summary</h2>')
        pdf.set_font('DejaVu', '', 12)
        pdf.write_html(f'<p>{cv_data["professional_summary"]}</p>')
        pdf.ln(5)

        # Add Skills
        pdf.set_font('DejaVu', 'B', 14)
        pdf.write_html('<h2>Skills</h2>')
        pdf.set_font('DejaVu', '', 12)
        pdf.write_html(f'<p>{", ".join(cv_data["skills"])}</p>')
        pdf.ln(5)

        # Add Key Achievements
        if cv_data.get('key_achievements'):
            pdf.set_font('DejaVu', 'B', 14)
            pdf.write_html('<h2>Key Achievements</h2>')
            pdf.set_font('DejaVu', '', 12)
            for achievement in cv_data['key_achievements']:
                pdf.write_html(f'<p>• {achievement}</p>')
            pdf.ln(5)

        # Add Experience
        pdf.set_font('DejaVu', 'B', 14)
        pdf.write_html('<h2>Experience</h2>')
        for job in cv_data['experience']:
            pdf.set_font('DejaVu', 'B', 12)
            pdf.write_html(
                f'<h3>{job["job_title"]} - {job["company"]}</h3>'
                f'<p>{job["date_range"]}</p>'
            )
            if job.get('location'):
                pdf.write_html(f'<p>{job["location"]}</p>')
            
            pdf.set_font('DejaVu', '', 12)
            for bullet in job['bullet_points']:
                pdf.write_html(f'<p style="margin-left: 10px;">• {bullet}</p>')
            pdf.ln(5)

        # Add Education
        pdf.set_font('DejaVu', 'B', 14)
        pdf.write_html('<h2>Education</h2>')
        pdf.set_font('DejaVu', '', 12)
        for edu in cv_data['education']:
            pdf.write_html(
                f'<p>{edu["degree"]} - {edu["institution"]} ({edu["graduation_year"]})</p>'
            )

        # Add Certifications
        if cv_data.get('certifications'):
            pdf.set_font('DejaVu', 'B', 14)
            pdf.write_html('<h2>Certifications</h2>')
            pdf.set_font('DejaVu', '', 12)
            for cert in cv_data['certifications']:
                pdf.write_html(
                    f'<p>{cert["name"]} ({cert["year"]}) - {cert["issuing_organization"]}</p>'
                )

        # Save the file
        pdf.output(output_path)
        return output_path 